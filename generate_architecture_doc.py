#!/usr/bin/env python3
"""
Генератор документа архитектуры ПО для платформы test-hard
Security Hardening & Monitoring Platform
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def set_cell_shading(cell, color):
    """Установить фоновый цвет ячейки"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, data, header_color='4472C4'):
    """Создать стилизованную таблицу"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = str(text)
            if i == 0:
                set_cell_shading(row.cells[j], header_color)
                run = row.cells[j].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    return table


def add_diagram(doc, diagram_text, font_size=7, caption=None):
    """Добавить ASCII диаграмму"""
    para = doc.add_paragraph()
    run = para.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    if caption:
        cap = doc.add_paragraph(caption)
        cap.italic = True


def create_title_page(doc):
    """Создать титульную страницу"""
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ДОКУМЕНТ АРХИТЕКТУРЫ\nПРОГРАММНОГО ОБЕСПЕЧЕНИЯ')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\n\ntest-hard\nПлатформа Security Hardening & Monitoring')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(68, 114, 196)
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f'\n\n\nВерсия: 2.0.0\nДата: {datetime.now().strftime("%d.%m.%Y")}')
    run.font.size = Pt(12)
    
    doc.add_page_break()


def create_toc(doc):
    """Создать оглавление"""
    doc.add_heading('Оглавление', level=1)
    toc = [
        '1. Введение',
        '2. Представление контекста (Context View)',
        '3. Логическое представление (Logical View)',
        '4. Представление вариантов использования (Use-Case View)',
        '5. Представление процессов (Process View)',
        '6. Физическое представление (Physical View)',
        '7. Представление данных (Data View)',
        '8. Операционное представление (Operational View)',
        '9. Представление информационной безопасности (Security View)',
        '10. Нефункциональные требования',
        '11. Прогноз по вычислительным мощностям',
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()


def create_introduction(doc):
    """Создать раздел введения"""
    doc.add_heading('1. Введение', level=1)
    
    doc.add_heading('1.1 Цель документа', level=2)
    doc.add_paragraph(
        'Настоящий документ описывает архитектуру платформы test-hard — '
        'автоматизированной системы для security hardening и мониторинга '
        'контейнеризированных и физических инфраструктур.'
    )
    
    doc.add_heading('1.2 Область применения', level=2)
    doc.add_paragraph(
        'Платформа test-hard предоставляет комплексное решение для:\n'
        '• Автоматического сканирования безопасности контейнеров (OpenSCAP, Lynis)\n'
        '• Мониторинга метрик безопасности (Prometheus, Grafana, Alertmanager)\n'
        '• Централизованного сбора логов (Loki, Promtail)\n'
        '• Тестирования техник MITRE ATT&CK (Atomic Red Team)\n'
        '• GitOps развертывания (ArgoCD)\n'
        '• Поддержки multi-environment (dev/staging/prod)'
    )
    
    doc.add_heading('1.3 Определения и сокращения', level=2)
    abbr_data = [
        ('Термин', 'Описание'),
        ('HLD', 'High-Level Design — высокоуровневый дизайн'),
        ('OpenSCAP', 'Security Content Automation Protocol'),
        ('Lynis', 'Инструмент аудита безопасности для Unix'),
        ('MITRE ATT&CK', 'Фреймворк описания тактик кибератак'),
        ('GitOps', 'Методология управления инфраструктурой через Git'),
        ('ArgoCD', 'Kubernetes-native инструмент непрерывной доставки'),
        ('Prometheus', 'Система мониторинга и алертинга'),
        ('Grafana', 'Платформа визуализации метрик'),
        ('Loki', 'Система агрегации логов'),
        ('K8s', 'Kubernetes — система оркестрации контейнеров'),
    ]
    add_styled_table(doc, abbr_data)
    doc.add_page_break()


def create_context_view(doc):
    """Создать представление контекста"""
    doc.add_heading('2. Представление контекста (Context View)', level=1)
    
    doc.add_heading('2.1 Описание', level=2)
    doc.add_paragraph(
        'Представление контекста показывает взаимодействие системы test-hard '
        'с внешними акторами и системами.'
    )
    
    doc.add_heading('2.2 Диаграмма контекста', level=2)
    context_diagram = '''
+-----------------------------------------------------------------------------------+
|                              ВНЕШНИЕ АКТОРЫ                                       |
|                                                                                   |
|  +------------+    +------------+    +------------+    +------------+             |
|  |  DevOps    |    |  Security  |    |   Admin    |    | Developer  |             |
|  |  Engineer  |    |  Analyst   |    |            |    |            |             |
|  +-----+------+    +-----+------+    +-----+------+    +-----+------+             |
|        |                 |                 |                 |                    |
|        +--------+--------+---------+-------+---------+-------+                    |
|                 |                  |                 |                            |
|                 v                  v                 v                            |
|  +------------------------------------------------------------------------+       |
|  |                      ПЛАТФОРМА test-hard                               |       |
|  |  +------------------------------------------------------------------+  |       |
|  |  |                     Web Dashboard (React + FastAPI)              |  |       |
|  |  +------------------------------------------------------------------+  |       |
|  |  +--------------+ +--------------+ +--------------+ +--------------+  |       |
|  |  |   OpenSCAP   | |    Lynis     | |   Atomic     | |   Telegraf   |  |       |
|  |  |   Scanner    | |   Scanner    | |  Red Team    | |    Agent     |  |       |
|  |  +--------------+ +--------------+ +--------------+ +--------------+  |       |
|  |  +------------------------------------------------------------------+  |       |
|  |  |         Monitoring Stack (Prometheus + Grafana + Loki)           |  |       |
|  |  +------------------------------------------------------------------+  |       |
|  +------------------------------------------------------------------------+       |
|                 |                  |                 |                            |
|        +--------+--------+---------+-------+---------+-------+                    |
|        v                 v                 v                 v                    |
|  +------------+    +------------+    +------------+    +------------+             |
|  |   Docker   |    | Kubernetes |    |   GitHub   |    |Notification|             |
|  |   Engine   |    |  Cluster   |    |   (CI/CD)  |    |  Systems   |             |
|  +------------+    +------------+    +------------+    +------------+             |
|                                                                                   |
|                           ВНЕШНИЕ СИСТЕМЫ                                         |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, context_diagram, 7, 'Рисунок 2.1 — Диаграмма контекста системы')
    
    doc.add_heading('2.3 Внешние акторы', level=2)
    actors = [
        ('Актор', 'Роль', 'Взаимодействие'),
        ('DevOps Engineer', 'Развертывание и обслуживание', 'CI/CD, Kubernetes, Docker'),
        ('Security Analyst', 'Анализ результатов', 'Grafana дашборды, отчёты'),
        ('Administrator', 'Управление конфигурацией', 'Web Dashboard, настройки'),
        ('Developer', 'Интеграция и расширение', 'API, скрипты'),
    ]
    add_styled_table(doc, actors)
    
    doc.add_heading('2.4 Внешние системы', level=2)
    systems = [
        ('Система', 'Назначение', 'Интерфейс'),
        ('Docker Engine', 'Контейнеризация', 'Docker API через Proxy'),
        ('Kubernetes', 'Оркестрация', 'kubectl, Kustomize, ArgoCD'),
        ('GitHub', 'CI/CD, хранение кода', 'Actions, Container Registry'),
        ('Notification Systems', 'Оповещения', 'Webhooks, SMTP'),
    ]
    add_styled_table(doc, systems)
    doc.add_page_break()


def create_logical_view(doc):
    """Создать логическое представление"""
    doc.add_heading('3. Логическое представление (Logical View)', level=1)
    
    doc.add_heading('3.1 High-Level Design (HLD)', level=2)
    hld_diagram = '''
+-----------------------------------------------------------------------------------+
|                          HIGH-LEVEL DESIGN (HLD)                                  |
|                          Платформа test-hard v2.0                                 |
+-----------------------------------------------------------------------------------+
|                                                                                   |
|  +------------------------ PRESENTATION LAYER -----------------------------+      |
|  |  +----------------+    +----------------+    +------------------+       |      |
|  |  | Web Dashboard  |    |    Grafana     |    |   API Gateway    |       |      |
|  |  | (React + Vite) |    | (Visualization)|    |   (FastAPI)      |       |      |
|  |  +----------------+    +----------------+    +------------------+       |      |
|  +-------------------------------------------------------------------------+      |
|                                     |                                             |
|  +------------------------ APPLICATION LAYER ------------------------------+      |
|  |  +--------------+ +--------------+ +--------------+ +-------------+    |      |
|  |  | Scan Service | |Schedule Svc  | | Report Svc   | | Alert Svc   |    |      |
|  |  | - OpenSCAP   | | - APScheduler| | - HTML/XML   | | - Rules     |    |      |
|  |  | - Lynis      | | - Cron       | | - Prometheus | | - Routing   |    |      |
|  |  | - Atomic RT  | | - Triggers   | | - JSON       | | - Actions   |    |      |
|  |  +--------------+ +--------------+ +--------------+ +-------------+    |      |
|  +-------------------------------------------------------------------------+      |
|                                     |                                             |
|  +------------------------ DATA LAYER -------------------------------------+      |
|  |  +--------------+ +--------------+ +--------------+ +-------------+    |      |
|  |  |  Prometheus  | |     Loki     | |   SQLite     | |   Volumes   |    |      |
|  |  |  (Metrics)   | |   (Logs)     | |  (App Data)  | |  (Reports)  |    |      |
|  |  +--------------+ +--------------+ +--------------+ +-------------+    |      |
|  +-------------------------------------------------------------------------+      |
|                                     |                                             |
|  +------------------------ INFRASTRUCTURE LAYER ---------------------------+      |
|  |  +--------------+ +--------------+ +--------------+ +-------------+    |      |
|  |  |   Docker     | |  Kubernetes  | |   ArgoCD     | |   GitHub    |    |      |
|  |  | - Compose    | | - Kustomize  | |  (GitOps)    | |  Actions    |    |      |
|  |  | - Socket Prx | | - NetworkPol | | - Auto-sync  | | - CI/CD     |    |      |
|  |  +--------------+ +--------------+ +--------------+ +-------------+    |      |
|  +-------------------------------------------------------------------------+      |
|                                                                                   |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, hld_diagram, 7, 'Рисунок 3.1 — High-Level Design диаграмма')
    
    doc.add_heading('3.2 Описание слоёв архитектуры', level=2)
    
    doc.add_heading('3.2.1 Presentation Layer', level=3)
    doc.add_paragraph(
        '• Web Dashboard — пользовательский интерфейс (React + TailwindCSS + Vite)\n'
        '• Grafana — визуализация метрик и дашборды\n'
        '• API Gateway — FastAPI с JWT аутентификацией и RBAC'
    )
    
    doc.add_heading('3.2.2 Application Layer', level=3)
    doc.add_paragraph(
        '• Scan Service — выполнение сканирований (OpenSCAP, Lynis, Atomic Red Team)\n'
        '• Schedule Service — планирование и автоматизация (APScheduler)\n'
        '• Report Service — генерация отчётов (HTML, XML, Prometheus metrics)\n'
        '• Alert Service — управление алертами (Alertmanager)'
    )
    
    doc.add_heading('3.2.3 Data Layer', level=3)
    doc.add_paragraph(
        '• Prometheus — хранение метрик (TSDB, retention 30d)\n'
        '• Loki — агрегация логов\n'
        '• SQLite — данные приложения (сканы, расписания, пользователи)\n'
        '• Volumes — файлы отчётов (HTML, XML, JSON)'
    )
    
    doc.add_heading('3.2.4 Infrastructure Layer', level=3)
    doc.add_paragraph(
        '• Docker — контейнеризация и Docker Compose\n'
        '• Kubernetes — оркестрация в production (Kustomize overlays)\n'
        '• ArgoCD — GitOps деплоймент\n'
        '• GitHub Actions — CI/CD пайплайны'
    )
    doc.add_page_break()


def create_usecase_view(doc):
    """Создать представление вариантов использования"""
    doc.add_heading('4. Представление вариантов использования (Use-Case View)', level=1)
    
    doc.add_heading('4.1 Диаграмма вариантов использования', level=2)
    usecase_diagram = '''
+-----------------------------------------------------------------------------------+
|                      USE-CASE DIAGRAM — test-hard                                 |
+-----------------------------------------------------------------------------------+
|                                                                                   |
|   +--------+                                              +--------+              |
|   | DevOps |                                              |Security|              |
|   |Engineer|                                              | Analyst|              |
|   +---+----+                                              +---+----+              |
|       |                                                       |                   |
|       |    +---------------------------------------------+    |                   |
|       |    |           ПЛАТФОРМА test-hard               |    |                   |
|       |    |                                             |    |                   |
|       +----+-> UC1: Развернуть платформу                 |    |                   |
|       |    |                                             |    |                   |
|       +----+-> UC2: Настроить мониторинг                 +----+                   |
|       |    |                                             |    |                   |
|       |    |   UC3: Запустить сканирование <-------------+----+                   |
|       |    |       +-> UC3.1: OpenSCAP scan              |    |                   |
|       |    |       +-> UC3.2: Lynis audit                |    |                   |
|       |    |       +-> UC3.3: Atomic Red Team            |    |                   |
|       |    |                                             |    |                   |
|       |    |   UC4: Просмотреть результаты <-------------+----+                   |
|       |    |       +-> UC4.1: Grafana dashboards         |    |                   |
|       |    |       +-> UC4.2: HTML/XML reports           |    |                   |
|       |    |                                             |    |                   |
|       |    |   UC5: Настроить расписание <---------------+----+                   |
|       |    |                                             |    |                   |
|       |    |   UC6: Получить алерты <--------------------+----+                   |
|       |    |                                             |    |                   |
|       +----+-> UC7: Управлять K8s deployment             |    |                   |
|       |    |                                             |    |                   |
|       +----+-> UC8: Обновить платформу (GitOps)          |    |                   |
|       |    |                                             |    |                   |
|       |    +---------------------------------------------+    |                   |
|       |                                                       |                   |
|   +---+----+                                              +---+----+              |
|   | Admin  |                                              |Developer              |
|   +--------+                                              +--------+              |
|                                                                                   |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, usecase_diagram, 7, 'Рисунок 4.1 — Диаграмма вариантов использования')
    
    doc.add_heading('4.2 Описание Use Cases', level=2)
    uc_data = [
        ('ID', 'Название', 'Описание'),
        ('UC1', 'Развернуть платформу', 'Развертывание через Docker Compose или Kubernetes'),
        ('UC2', 'Настроить мониторинг', 'Конфигурация Prometheus, Grafana, Alertmanager'),
        ('UC3', 'Запустить сканирование', 'Выполнение security scans (OpenSCAP, Lynis, ART)'),
        ('UC4', 'Просмотреть результаты', 'Анализ через дашборды и отчёты'),
        ('UC5', 'Настроить расписание', 'Автоматизация сканирований по cron'),
        ('UC6', 'Получить алерты', 'Получение уведомлений о проблемах'),
        ('UC7', 'Управлять K8s', 'Управление через kubectl, ArgoCD'),
        ('UC8', 'Обновить платформу', 'GitOps обновление через PR'),
    ]
    add_styled_table(doc, uc_data)
    doc.add_page_break()


def create_process_view(doc):
    """Создать представление процессов"""
    doc.add_heading('5. Представление процессов (Process View)', level=1)
    
    doc.add_heading('5.1 Процесс сканирования безопасности', level=2)
    scan_process = '''
+-----------------------------------------------------------------------------------+
|                    SEQUENCE DIAGRAM: Security Scanning Process                    |
+-----------------------------------------------------------------------------------+
|                                                                                   |
|   User          Dashboard       Scanner        DockerProxy      Target            |
|    |                |              |               |              |               |
|    |  1. Start Scan |              |               |              |               |
|    |--------------->|              |               |              |               |
|    |                | 2. API Call  |               |              |               |
|    |                |------------->|               |              |               |
|    |                |              | 3. Get Containers            |               |
|    |                |              |-------------->|              |               |
|    |                |              | 4. Container IDs             |               |
|    |                |              |<--------------|              |               |
|    |                |              | 5. Exec Scan  |              |               |
|    |                |              |-------------->|------------->|               |
|    |                |              |               |  6. Execute  |               |
|    |                |              |               |  OpenSCAP/   |               |
|    |                |              |               |  Lynis       |               |
|    |                |              |               |<-------------|               |
|    |                |              | 7. Scan Results              |               |
|    |                |              |<--------------|              |               |
|    |                | 8. Parse Results             |              |               |
|    |                |<-------------|               |              |               |
|    |                | 9. Store Metrics (Prometheus)|              |               |
|    | 10. Status OK  |              |               |              |               |
|    |<---------------|              |               |              |               |
|                                                                                   |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, scan_process, 7, 'Рисунок 5.1 — Sequence диаграмма сканирования')
    
    doc.add_heading('5.2 Потоки данных мониторинга', level=2)
    monitoring_flow = '''
+-----------------------------------------------------------------------------------+
|                     DATA FLOW: Monitoring & Alerting                              |
+-----------------------------------------------------------------------------------+
|                                                                                   |
|  +-------------+                                                                  |
|  |  Targets    |                                                                  |
|  | (Scanners)  |                                                                  |
|  +------+------+                                                                  |
|         |                                                                         |
|         | Metrics (pull every 15s)                                                |
|         v                                                                         |
|  +-------------+         +-------------+         +-------------+                  |
|  |  Telegraf   |-------->| Prometheus  |-------->| Alertmanager|                  |
|  | (Collector) | scrape  |   (TSDB)    | alerts  |  (Routing)  |                  |
|  +-------------+         +------+------+         +------+------+                  |
|                                 |                       |                         |
|         +--------- ------------+                       |                         |
|         |                      |                       v                         |
|         v                      v                +-------------+                  |
|  +-------------+         +-------------+        |Notifications|                  |
|  |   Grafana   |<--------|    Loki     |        | (Webhook,   |                  |
|  | (Dashboards)| query   |   (Logs)    |        |   Email)    |                  |
|  +-------------+         +------+------+        +-------------+                  |
|                                 ^                                                 |
|                                 |                                                 |
|                          +------+------+                                          |
|                          |  Promtail   |                                          |
|                          | (Log ship)  |                                          |
|                          +-------------+                                          |
|                                 ^                                                 |
|                          +------+------+                                          |
|                          | Container   |                                          |
|                          |   Logs      |                                          |
|                          +-------------+                                          |
|                                                                                   |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, monitoring_flow, 7, 'Рисунок 5.2 — Data Flow диаграмма мониторинга')
    doc.add_page_break()


def create_physical_view(doc):
    """Создать физическое представление"""
    doc.add_heading('6. Физическое представление (Physical View)', level=1)
    
    doc.add_heading('6.1 Deployment Diagram — Docker Compose', level=2)
    docker_deploy = '''
+-----------------------------------------------------------------------------------+
|                  DEPLOYMENT DIAGRAM: Docker Compose Environment                   |
+-----------------------------------------------------------------------------------+
|                                                                                   |
|  +--------------------------------- HOST -------------------------------------+   |
|  |                                                                            |   |
|  |  +----------------------- Network: default ----------------------------+  |   |
|  |  |  +------------+ +------------+ +------------+ +------------+        |  |   |
|  |  |  | prometheus | |  grafana   | |    loki    | |  promtail  |        |  |   |
|  |  |  |   :9090    | |   :3000    | |   :3100    | |            |        |  |   |
|  |  |  | 1CPU/1GB   | | 0.5CPU/512M| | 0.5CPU/512M| | 0.25CPU/256M       |  |   |
|  |  |  +-----+------+ +-----+------+ +-----+------+ +------------+        |  |   |
|  |  |        |              |              |                              |  |   |
|  |  |  +-----+------+ +-----+------+ +-----+------+                       |  |   |
|  |  |  |alertmanager| |  telegraf  | | openscap-  |                       |  |   |
|  |  |  |   :9093    | |   :9091    | |  scanner   |                       |  |   |
|  |  |  | 0.25CPU/128M| 0.5CPU/256M | | 1CPU/512M  |                       |  |   |
|  |  |  +------------+ +------------+ +-----+------+                       |  |   |
|  |  |                                      |                              |  |   |
|  |  |  +------------+                      |                              |  |   |
|  |  |  |lynis-scan  |                      |                              |  |   |
|  |  |  | 1CPU/512M  |                      |                              |  |   |
|  |  |  +-----+------+                      |                              |  |   |
|  |  +--------+-----------------------------+------------------------------+  |   |
|  |           |                             |                                 |   |
|  |  +--------+---- Network: scanner-net ---+-----------------------------+   |   |
|  |  |        |    +------------------+     |                             |   |   |
|  |  |        +--->| docker-proxy     |<----+                             |   |   |
|  |  |             |    :2375         |                                   |   |   |
|  |  |             | 0.25CPU/128M     |                                   |   |   |
|  |  |             +--------+---------+                                   |   |   |
|  |  +----------------------+---------------------------------------------+   |   |
|  |                         |                                                 |   |
|  |  +----------------------+---------------------------------------------+   |   |
|  |  |                      v                                             |   |   |
|  |  |  +-------------------------------------------------------------+   |   |   |
|  |  |  |              TARGET CONTAINERS (sleep infinity)             |   |   |   |
|  |  |  |  +-------+ +-------+ +-------+ +-------+ +-------+          |   |   |   |
|  |  |  |  |fedora | |debian | |centos | |ubuntu | |  alt  |          |   |   |   |
|  |  |  |  |  :40  | |  :12  | |strm:9 | | :22.04| | linux |          |   |   |   |
|  |  |  |  +-------+ +-------+ +-------+ +-------+ +-------+          |   |   |   |
|  |  |  +-------------------------------------------------------------+   |   |   |
|  |  +--------------------------------------------------------------------+   |   |
|  |                                                                           |   |
|  |  VOLUMES: prometheus-data, grafana-data, loki-data, ./reports             |   |
|  +--------------------------------------------------------------------------+|   |
|                                                                               |   |
|  PORTS: 9090 (Prometheus), 3000 (Grafana), 9093 (Alertmanager), 3100 (Loki)  |   |
|                                                                               |   |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, docker_deploy, 6, 'Рисунок 6.1 — Deployment диаграмма Docker Compose')
    doc.add_page_break()
    
    doc.add_heading('6.2 Deployment Diagram — Kubernetes', level=2)
    k8s_deploy = '''
+-----------------------------------------------------------------------------------+
|                  DEPLOYMENT DIAGRAM: Kubernetes Environment                       |
+-----------------------------------------------------------------------------------+
|                                                                                   |
|  +--------------------------- KUBERNETES CLUSTER ----------------------------+   |
|  |                                                                            |   |
|  |  +---------------------- Namespace: monitoring ------------------------+  |   |
|  |  |                                                                     |  |   |
|  |  |  +------------------- Network Policies ---------------------------+ |  |   |
|  |  |  |  * default-deny-all   (Deny all by default)                    | |  |   |
|  |  |  |  * allow-grafana      (Ingress + Prometheus egress)            | |  |   |
|  |  |  |  * allow-prometheus   (Telegraf scrape + Alertmanager egress)  | |  |   |
|  |  |  |  * allow-telegraf     (Prometheus scrape ingress)              | |  |   |
|  |  |  |  * allow-alertmanager (Prometheus ingress + external egress)   | |  |   |
|  |  |  +----------------------------------------------------------------+ |  |   |
|  |  |                                                                     |  |   |
|  |  |  +----------+ +----------+ +----------+ +----------+               |  |   |
|  |  |  |Deployment| |Deployment| |Deployment| |Deployment|               |  |   |
|  |  |  |prometheus| | grafana  | | telegraf | |alertmgr  |               |  |   |
|  |  |  |replicas:1| |replicas:1| |replicas:1| |replicas:1|               |  |   |
|  |  |  +----+-----+ +----+-----+ +----+-----+ +----+-----+               |  |   |
|  |  |       |            |            |            |                      |  |   |
|  |  |  +----+-----+ +----+-----+ +----+-----+ +----+-----+               |  |   |
|  |  |  | Service  | | Service  | | Service  | | Service  |               |  |   |
|  |  |  |  :9090   | |  :3000   | |  :9091   | |  :9093   |               |  |   |
|  |  |  |ClusterIP | |ClusterIP | |ClusterIP | |ClusterIP |               |  |   |
|  |  |  +----------+ +----------+ +----------+ +----------+               |  |   |
|  |  |                                                                     |  |   |
|  |  |  +---------------------- ConfigMaps -----------------------------+ |  |   |
|  |  |  |  * prometheus-config    (prometheus.yml, alert.rules.yml)     | |  |   |
|  |  |  |  * grafana-datasources  (datasource configuration)            | |  |   |
|  |  |  |  * telegraf-config      (telegraf.conf)                       | |  |   |
|  |  |  +----------------------------------------------------------------+ |  |   |
|  |  |                                                                     |  |   |
|  |  |  +-------------- PersistentVolumeClaims -------------------------+ |  |   |
|  |  |  |  * prometheus-data-pvc  (10Gi)                                 | |  |   |
|  |  |  |  * grafana-data-pvc     (5Gi)                                  | |  |   |
|  |  |  +----------------------------------------------------------------+ |  |   |
|  |  |                                                                     |  |   |
|  |  +---------------------------------------------------------------------+  |   |
|  |                                                                            |   |
|  |  +---------------------------- ArgoCD ---------------------------------+  |   |
|  |  |  * application-dev.yaml     (auto-sync enabled)                     |  |   |
|  |  |  * application-staging.yaml (auto-sync enabled)                     |  |   |
|  |  |  * application-prod.yaml    (manual sync, auto-prune disabled)      |  |   |
|  |  +--------------------------------------------------------------------|  |   |
|  |                                                                            |   |
|  +----------------------------------------------------------------------------+   |
|                                                                                   |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, k8s_deploy, 6, 'Рисунок 6.2 — Deployment диаграмма Kubernetes')
    doc.add_page_break()


def create_data_view(doc):
    """Создать представление данных"""
    doc.add_heading('7. Представление данных (Data View)', level=1)
    
    doc.add_heading('7.1 Модель данных', level=2)
    doc.add_paragraph(
        'Система использует несколько хранилищ данных для различных целей:'
    )
    
    data_stores = [
        ('Хранилище', 'Технология', 'Назначение', 'Retention'),
        ('Метрики', 'Prometheus TSDB', 'Временные ряды метрик', '30 дней / 10GB'),
        ('Логи', 'Loki', 'Агрегированные логи', 'Настраиваемый'),
        ('App Data', 'SQLite', 'Пользователи, расписания, сканы', 'Постоянно'),
        ('Отчёты', 'File Volumes', 'HTML/XML/JSON отчёты', 'Постоянно'),
        ('Config', 'ConfigMaps/Files', 'Конфигурация сервисов', 'Версионировано'),
    ]
    add_styled_table(doc, data_stores)
    
    doc.add_heading('7.2 Схема базы данных Dashboard', level=2)
    db_schema = '''
+-----------------------------------------------------------------------------------+
|                       DATABASE SCHEMA: Dashboard (SQLite)                         |
+-----------------------------------------------------------------------------------+
|                                                                                   |
|  +------------------+       +------------------+       +------------------+       |
|  |      users       |       |      scans       |       |    schedules     |       |
|  +------------------+       +------------------+       +------------------+       |
|  | PK id: INTEGER   |       | PK id: INTEGER   |       | PK id: INTEGER   |       |
|  | username: TEXT   |       | target_id: TEXT  |       | name: TEXT       |       |
|  | email: TEXT      |       | scan_type: TEXT  |       | cron_expr: TEXT  |       |
|  | hashed_pwd: TEXT |       | status: TEXT     |       | scan_type: TEXT  |       |
|  | role: TEXT       |       | started_at: DT   |       | targets: JSON    |       |
|  | is_active: BOOL  |<---+  | completed_at: DT |       | enabled: BOOL    |       |
|  | created_at: DT   |    |  | result_path: TEXT|       | created_at: DT   |       |
|  +------------------+    |  | created_by: FK   |-------| last_run: DT     |       |
|                          |  +------------------+       +------------------+       |
|                          |                                                        |
|                          +--------------------------------------------------------+
|                                                                                   |
|  ROLES: viewer (read-only), operator (run scans), admin (full access)            |
|                                                                                   |
+-----------------------------------------------------------------------------------+
'''
    add_diagram(doc, db_schema, 7, 'Рисунок 7.1 — Схема базы данных Dashboard')
    
    doc.add_heading('7.3 Метрики безопасности', level=2)
    metrics = [
        ('Метрика', 'Источник', 'Описание'),
        ('lynis_score', 'Lynis', 'Общий индекс hardening (0-100)'),
        ('lynis_warnings_count', 'Lynis', 'Количество предупреждений'),
        ('openscap_pass_count', 'OpenSCAP', 'Количество пройденных правил'),
        ('openscap_fail_count', 'OpenSCAP', 'Количество проваленных правил'),
        ('art_test_result', 'Atomic RT', 'Результат атомарного теста'),
        ('art_scenario_status', 'Atomic RT', 'Статус сценария MITRE ATT&CK'),
    ]
    add_styled_table(doc, metrics)
    doc.add_page_break()


def create_operational_view(doc):
    """Создать операционное представление"""
    doc.add_heading('8. Операционное представление (Operational View)', level=1)
    
    doc.add_heading('8.1 Развертывание', level=2)
    doc.add_paragraph(
        'Платформа поддерживает несколько режимов развертывания:'
    )
    
    deploy_modes = [
        ('Режим', 'Команда', 'Описание'),
        ('Full Stack', 'make up', 'Все компоненты + мониторинг + сканеры'),
        ('With Logging', 'make up-with-logging', 'Full Stack + Loki + Promtail'),
        ('Targets Only', 'make up-targets', 'Только целевые контейнеры'),
        ('Monitor Only', 'make monitor', 'Prometheus + Grafana + Alertmanager'),
        ('Scan', 'make scan', 'Запуск сканирования'),
    ]
    add_styled_table(doc, deploy_modes)
    
    doc.add_heading('8.2 Мониторинг состояния', level=2)
    
    health_checks = [
        ('Сервис', 'Endpoint', 'Интервал'),
        ('Prometheus', 'GET /-/healthy', '30s'),
        ('Grafana', 'GET /api/health', '30s'),
        ('Telegraf', 'GET /metrics', '30s'),
        ('Alertmanager', 'GET /-/healthy', '10s'),
        ('Loki', 'GET /ready', '30s'),
    ]
    add_styled_table(doc, health_checks)
    
    doc.add_heading('8.3 CI/CD Pipeline', level=2)
    doc.add_paragraph(
        'GitHub Actions workflows:\n'
        '• Security Scanning — автоматическое сканирование при push\n'
        '• Docker Build — проверка сборки образов\n'
        '• Pytest Tests — unit и integration тесты\n'
        '• Secret Scanning — поиск случайно закоммиченных секретов\n'
        '• Container Registry — публикация в GHCR с Cosign signing'
    )
    doc.add_page_break()


def create_security_view(doc):
    """Создать представление информационной безопасности"""
    doc.add_heading('9. Представление информационной безопасности (Security View)', level=1)
    
    doc.add_heading('9.1 Модель угроз', level=2)
    doc.add_paragraph(
        'Основные векторы атак и меры защиты:'
    )
    
    threats = [
        ('Угроза', 'Риск', 'Мера защиты'),
        ('Доступ к Docker socket', 'Высокий', 'Docker Socket Proxy с ограничениями'),
        ('Неавторизованный доступ', 'Высокий', 'JWT аутентификация, RBAC'),
        ('Утечка секретов', 'Высокий', '.env файлы, Secret Scanning в CI'),
        ('Эскалация привилегий', 'Средний', 'no-new-privileges, минимальные capabilities'),
        ('Сетевые атаки', 'Средний', 'Network Policies, изоляция сетей'),
        ('DoS атаки', 'Средний', 'Resource limits, rate limiting'),
    ]
    add_styled_table(doc, threats)
    
    doc.add_heading('9.2 Меры безопасности', level=2)
    
    doc.add_heading('9.2.1 Docker Socket Proxy', level=3)
    doc.add_paragraph(
        '• Используется tecnativa/docker-socket-proxy\n'
        '• Ограниченные API endpoints: CONTAINERS, EXEC, IMAGES, INFO\n'
        '• Доступ только для чтения где возможно\n'
        '• Нет доступа к POST операциям с VOLUMES, NETWORKS'
    )
    
    doc.add_heading('9.2.2 Container Security', level=3)
    doc.add_paragraph(
        '• cap_drop: ALL — сброс всех capabilities\n'
        '• cap_add: только необходимые (SYS_PTRACE, SYS_ADMIN, NET_ADMIN)\n'
        '• security_opt: no-new-privileges:true\n'
        '• Resource limits для всех контейнеров'
    )
    
    doc.add_heading('9.2.3 Network Security', level=3)
    doc.add_paragraph(
        '• Изоляция сетей: default, scanner-net\n'
        '• Kubernetes Network Policies: deny-all по умолчанию\n'
        '• Явные allow правила для необходимого трафика'
    )
    
    doc.add_heading('9.3 Управление пользователями', level=2)
    roles = [
        ('Роль', 'Права', 'Описание'),
        ('hardening-admin', 'Ограниченный sudo', 'Администрирование платформы'),
        ('hardening-scanner', 'Минимальные права', 'Запуск сканирований'),
        ('hardening-service', 'Без sudo', 'Автоматизация через systemd'),
        ('hardening-readonly', 'Только чтение', 'Просмотр отчётов'),
    ]
    add_styled_table(doc, roles)
    doc.add_page_break()


def create_nfr_view(doc):
    """Создать представление нефункциональных требований"""
    doc.add_heading('10. Нефункциональные требования', level=1)
    
    doc.add_heading('10.1 Производительность', level=2)
    perf_reqs = [
        ('Требование', 'Значение', 'Примечание'),
        ('Время сканирования', '< 5 минут', 'Для одного контейнера'),
        ('Scrape interval', '15 секунд', 'Prometheus'),
        ('Retention метрик', '30 дней или 10GB', 'Настраиваемо'),
        ('Test coverage', '> 80%', 'Unit + Integration'),
    ]
    add_styled_table(doc, perf_reqs)
    
    doc.add_heading('10.2 Масштабируемость', level=2)
    doc.add_paragraph(
        '• Поддержка multi-distribution: Debian, Ubuntu, Fedora, CentOS, ALT Linux\n'
        '• Kubernetes deployment с Kustomize overlays (dev/staging/prod)\n'
        '• ArgoCD для GitOps автоматизации\n'
        '• Horizontal scaling через replica count'
    )
    
    doc.add_heading('10.3 Доступность', level=2)
    doc.add_paragraph(
        '• Health checks для всех критических сервисов\n'
        '• Auto-restart политики в Docker и Kubernetes\n'
        '• Persistent volumes для данных'
    )
    
    doc.add_heading('10.4 Совместимость', level=2)
    compat = [
        ('Компонент', 'Версия', 'Примечание'),
        ('Python', '3.11+', 'Runtime'),
        ('Docker', '24.0+', 'Container runtime'),
        ('Kubernetes', '1.28+', 'Orchestration'),
        ('Prometheus', 'v2.52.0', 'Monitoring'),
        ('Grafana', '11.0.0', 'Visualization'),
        ('Loki', '2.9.0', 'Logging'),
    ]
    add_styled_table(doc, compat)
    doc.add_page_break()


def create_capacity_forecast(doc):
    """Создать прогноз по вычислительным мощностям"""
    doc.add_heading('11. Прогноз по вычислительным мощностям', level=1)
    
    doc.add_heading('11.1 Требования к ресурсам', level=2)
    resources = [
        ('Компонент', 'CPU', 'Memory', 'Storage'),
        ('Prometheus', '1 CPU', '1 GB', '10 GB (TSDB)'),
        ('Grafana', '0.5 CPU', '512 MB', '5 GB'),
        ('Loki', '0.5 CPU', '512 MB', '10 GB'),
        ('Promtail', '0.25 CPU', '256 MB', '-'),
        ('Telegraf', '0.5 CPU', '256 MB', '-'),
        ('Alertmanager', '0.25 CPU', '128 MB', '100 MB'),
        ('OpenSCAP Scanner', '1 CPU', '512 MB', '-'),
        ('Lynis Scanner', '1 CPU', '512 MB', '-'),
        ('Docker Proxy', '0.25 CPU', '128 MB', '-'),
        ('ИТОГО (минимум)', '5.25 CPU', '3.8 GB', '25 GB'),
    ]
    add_styled_table(doc, resources)
    
    doc.add_heading('11.2 Рекомендации по масштабированию', level=2)
    scaling = [
        ('Нагрузка', 'CPU', 'Memory', 'Storage'),
        ('Dev (до 10 targets)', '4 CPU', '8 GB', '50 GB'),
        ('Staging (до 50 targets)', '8 CPU', '16 GB', '100 GB'),
        ('Production (до 200 targets)', '16 CPU', '32 GB', '500 GB'),
    ]
    add_styled_table(doc, scaling)
    
    doc.add_heading('11.3 KPIs и метрики', level=2)
    kpis = [
        ('Метрика', 'Текущее', 'Цель Q2 2026', 'Цель Q4 2026'),
        ('Test Coverage', '80%', '85%', '90%'),
        ('Mean Time to Scan', '~5 min', '~3 min', '~2 min'),
        ('Supported Distributions', '5', '8', '12'),
        ('Documentation Coverage', '90%', '95%', '98%'),
    ]
    add_styled_table(doc, kpis)


def main():
    """Главная функция"""
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Создание разделов
    create_title_page(doc)
    create_toc(doc)
    create_introduction(doc)
    create_context_view(doc)
    create_logical_view(doc)
    create_usecase_view(doc)
    create_process_view(doc)
    create_physical_view(doc)
    create_data_view(doc)
    create_operational_view(doc)
    create_security_view(doc)
    create_nfr_view(doc)
    create_capacity_forecast(doc)
    
    # Сохранение документа
    output_path = 'Документ_Архитектуры_ПО_test-hard_v2.docx'
    doc.save(output_path)
    print(f'Документ сохранён: {output_path}')
    return output_path


if __name__ == '__main__':
    main()
